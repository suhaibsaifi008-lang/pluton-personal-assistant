"""
File operations — Word, Excel, and general files.
Every operation is restricted to config.ALLOWED_FOLDERS so Pluton can only
touch files you've explicitly opted in.
"""

import os
import glob

from docx import Document
from openpyxl import Workbook, load_workbook

import config


class PermissionDenied(Exception):
    pass


class FileManager:
    def __init__(self):
        self.allowed_folders = [
            os.path.abspath(os.path.expandvars(f)) for f in config.ALLOWED_FOLDERS
        ]
        self.default_output = os.path.expandvars(config.DEFAULT_OUTPUT_FOLDER)
        os.makedirs(self.default_output, exist_ok=True)

    def _check_allowed(self, path):
        path = os.path.abspath(os.path.expandvars(path))
        if not any(path.startswith(folder) for folder in self.allowed_folders):
            raise PermissionDenied(
                f"'{path}' is outside the folders you've allowed. "
                f"Add it to ALLOWED_FOLDERS in config.py first."
            )
        return path

    def find_file(self, name_hint, extensions=("docx", "xlsx")):
        """Search allowed folders for a file whose name contains name_hint."""
        matches = []
        for folder in self.allowed_folders:
            if not os.path.isdir(folder):
                continue
            for ext in extensions:
                pattern = os.path.join(folder, "**", f"*{name_hint}*.{ext}")
                matches.extend(glob.glob(pattern, recursive=True))
        return matches

    # ---------------- Word ----------------

    def read_docx(self, path):
        path = self._check_allowed(path)
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def append_to_docx(self, path, text, as_heading=False):
        path = self._check_allowed(path)
        doc = Document(path) if os.path.exists(path) else Document()
        if as_heading:
            doc.add_heading(text, level=1)
        else:
            for para in text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(path)

    def create_docx(self, path, title, sections):
        """sections: list of (heading, body) tuples. Body can be any length —
        each paragraph is added separately so there's no size limit."""
        path = self._check_allowed(path)
        doc = Document()
        if title:
            doc.add_heading(title, level=0)
        for heading, body in sections:
            if heading:
                doc.add_heading(heading, level=1)
            for para in body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(path)
        return path

    # ---------------- Excel ----------------

    def read_xlsx(self, path, max_rows=200):
        path = self._check_allowed(path)
        wb = load_workbook(path, data_only=True)
        ws = wb.active
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            rows.append(row)
        return rows

    def append_xlsx_row(self, path, row_values):
        path = self._check_allowed(path)
        wb = load_workbook(path) if os.path.exists(path) else Workbook()
        ws = wb.active
        ws.append(row_values)
        wb.save(path)

    def create_xlsx(self, path, headers, rows):
        """rows: list of lists. No row-count limit — openpyxl streams to disk."""
        path = self._check_allowed(path)
        wb = Workbook()
        ws = wb.active
        if headers:
            ws.append(headers)
        for row in rows:
            ws.append(row)
        wb.save(path)
        return path

    # ---------------- Plain text ----------------

    def read_text(self, path):
        path = self._check_allowed(path)
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def write_text(self, path, content, append=False):
        path = self._check_allowed(path)
        mode = "a" if append else "w"
        with open(path, mode, encoding="utf-8") as f:
            f.write(content)
        return path
